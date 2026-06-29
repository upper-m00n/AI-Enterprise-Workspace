from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.orm import Session
from typing import List
from uuid import UUID
import io
from pypdf import PdfReader
from docx import Document as DocxDocument

from app.api import deps
from app.db.database import get_db
from app.models.user import User
from app.models.document import Document, DocumentChunk
from app.schemas.document import DocumentResponse
from app.utils.text_splitter import SimpleTextSplitter
from app.embeddings.service import get_embeddings

router = APIRouter()

def extract_text_from_pdf(content_bytes: bytes) -> str:
    pdf_file = io.BytesIO(content_bytes)
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(content_bytes: bytes) -> str:
    docx_file = io.BytesIO(content_bytes)
    doc = DocxDocument(docx_file)
    text = ""
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"
    return text

@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Document:
    ext = file.filename.split(".")[-1].lower() if file.filename else ""
    allowed_exts = ["txt", "md", "pdf", "docx"]
    
    if ext not in allowed_exts:
        raise HTTPException(
            status_code=400,
            detail="Invalid file format. Supported formats: .txt, .md, .pdf, .docx"
        )

    try:
        content_bytes = await file.read()
        if ext == "pdf":
            content = extract_text_from_pdf(content_bytes)
        elif ext == "docx":
            content = extract_text_from_docx(content_bytes)
        else:
            content = content_bytes.decode("utf-8")
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Could not parse file content: {str(e)}"
        )

    if not content.strip():
        raise HTTPException(status_code=400, detail="The extracted document content is empty.")

    db_document = Document(
        filename=file.filename or f"unknown.{ext}",
        mime_type=file.content_type or "text/plain",
        file_size=len(content_bytes),
        user_id=current_user.id
    )
    db.add(db_document)
    db.commit()
    db.refresh(db_document)

    splitter = SimpleTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_text(content)

    if chunks:
        embeddings = get_embeddings(chunks)

        db_chunks = []
        for text, vector in zip(chunks, embeddings):
            chunk = DocumentChunk(
                document_id=db_document.id,
                content=text,
                embedding=vector
            )
            db_chunks.append(chunk)
        
        db.bulk_save_objects(db_chunks)
        db.commit()

    return db_document

@router.get("/", response_model=List[DocumentResponse])
def list_documents(
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_user)
) -> List[Document]:
    return db.query(Document).filter(Document.user_id == current_user.id).all()

@router.get("/{document_id}", response_model=DocumentResponse)
def get_document(
    document_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_user)
) -> Document:
    doc = db.query(Document).filter(
        Document.id == document_id, 
        Document.user_id == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc

@router.delete("/{document_id}")
def delete_document(
    document_id: UUID,
    db: Session = Depends(get_db),
    current_user: User = Depends(deps.get_current_user)
):
    doc = db.query(Document).filter(
        Document.id == document_id, 
        Document.user_id == current_user.id
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    db.delete(doc)
    db.commit()
    return {"message": "Document and associated chunks deleted successfully"}
